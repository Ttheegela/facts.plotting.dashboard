"""
make_docs.py — Convert FACTS reports from Markdown to .docx
Run: python make_docs.py
Outputs: amarel_singularity_report.docx, dashboard_size_optimization_report.docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPORTS_DIR = Path(__file__).parent


# ──────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────

def add_heading(doc, text, level):
    p = doc.add_heading(text.strip(), level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_code_block(doc, code):
    """Grey shaded paragraph for code."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table_row(tbl, cells, bold=False):
    row = tbl.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = text.strip()
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row


def apply_inline(run_text, paragraph):
    """Write text with **bold** markers into a paragraph."""
    parts = re.split(r"\*\*(.*?)\*\*", run_text)
    for j, part in enumerate(parts):
        r = paragraph.add_run(part)
        if j % 2 == 1:
            r.bold = True
    return paragraph


def add_para(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    apply_inline(text.strip(), p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    apply_inline(text.strip(), p)
    return p


def parse_md_table(lines):
    """Return (headers, rows) from markdown table lines."""
    headers, rows = [], []
    for i, ln in enumerate(lines):
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if i == 0:
            headers = cells
        elif re.match(r"^[\s|:\-]+$", ln):
            continue
        else:
            rows.append(cells)
    return headers, rows


def add_md_table(doc, headers, rows):
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = "Table Grid"
    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        if i < n_cols:
            cell = hdr_row.cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
    # Data rows
    for row_cells in rows:
        row = tbl.add_row()
        for i, text in enumerate(row_cells):
            if i < n_cols:
                row.cells[i].text = text
    doc.add_paragraph()  # spacing after table


# ──────────────────────────────────────────────────────────────
# Main parser
# ──────────────────────────────────────────────────────────────

def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()

    # Compact page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i]

        # Headings
        m = re.match(r"^(#{1,3})\s+(.*)", ln)
        if m:
            add_heading(doc, m.group(2), level=len(m.group(1)))
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", ln.strip()):
            i += 1
            continue

        # Code block (fenced)
        if ln.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1  # skip closing ```
            continue

        # Markdown table
        if ln.strip().startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(tbl_lines)
            if headers:
                add_md_table(doc, headers, rows)
            continue

        # Bullet
        m = re.match(r"^[-*]\s+(.*)", ln)
        if m:
            add_bullet(doc, m.group(1))
            i += 1
            continue

        # Blockquote
        m = re.match(r"^>\s*(.*)", ln)
        if m:
            add_para(doc, m.group(1), indent=True)
            i += 1
            continue

        # Empty line
        if not ln.strip():
            i += 1
            continue

        # Regular paragraph
        add_para(doc, ln)
        i += 1

    doc.save(str(docx_path))
    print(f"  Saved: {docx_path.name}")


if __name__ == "__main__":
    print("Generating Word documents...")
    md_to_docx(
        REPORTS_DIR / "amarel_singularity_report.md",
        REPORTS_DIR / "amarel_singularity_report.docx",
    )
    md_to_docx(
        REPORTS_DIR / "dashboard_size_optimization_report.md",
        REPORTS_DIR / "dashboard_size_optimization_report.docx",
    )
    print("Done.")
